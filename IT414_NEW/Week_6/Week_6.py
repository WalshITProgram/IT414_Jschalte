import os
import zipfile
import pandas as pd
import PyPDF4 
from docx import Document

def extract_raw_data(zip_file_path: str, extract_folder: str):
    """
    Extracts raw report data from a zip file.

    Args:
        zip_file_path (str): Path to the zip file.
        extract_folder (str): Path to the folder where the data will be extracted.
    """
    with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
        zip_ref.extractall(extract_folder)

def read_sales_data(file_path: str) -> pd.DataFrame:
    """
    Reads sales data from an Excel file.

    Args:
        file_path (str): Path to the Excel file.

    Returns:
        pd.DataFrame: Sales data.
    """
    return pd.read_excel(file_path)

def read_operations_data(file_path: str) -> list:
    """
    Reads operations data from a Word file.

    Args:
        file_path (str): Path to the Word file.

    Returns:
        list: List of paragraphs containing operations data.
    """
    doc = Document(file_path)
    return [para.text for para in doc.paragraphs]

def read_marketing_data(file_path: str) -> str:
    """
    Reads marketing data from a PDF file.

    Args:
        file_path (str): Path to the PDF file.

    Returns:
        str: Marketing data.
    """
    marketing_text = ''
    with open(file_path, 'rb') as pdf_file:
        pdf_reader = PyPDF4.PdfFileReader(pdf_file)
        for page_num in range(pdf_reader.numPages):
            page = pdf_reader.getPage(page_num)
            marketing_text += page.extractText()
    return marketing_text

def read_it_data(file_path: str) -> str:
    """
    Reads IT data from an HTML file.

    Args:
        file_path (str): Path to the HTML file.

    Returns:
        str: IT data.
    """
    with open(file_path, 'r') as html_file:
        return html_file.read()

def generate_ceo_report(sales_data: pd.DataFrame, operations_content: list, marketing_text: str, it_content: str, output_path: str):
    """
    Generates a CEO report.

    Args:
        sales_data (pd.DataFrame): Sales data.
        operations_content (list): List of paragraphs containing operations data.
        marketing_text (str): Marketing data.
        it_content (str): IT data.
        output_path (str): Path to save the CEO report.
    """
    doc = Document()
    doc.add_heading('CEO Report', level=1)

    # Add Sales section
    doc.add_heading('Sales', level=2)
    sales_table = doc.add_table(rows=len(sales_data) + 1, cols=len(sales_data.columns))
    for i, column in enumerate(sales_data.columns):
        sales_table.cell(0, i).text = column
    for i, row in enumerate(sales_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            sales_table.cell(i, j).text = str(value)

    # Add Operations section
    doc.add_page_break()
    doc.add_heading('Operations', level=2)
    for para in operations_content:
        doc.add_paragraph(para)

    # Add Marketing section
    doc.add_page_break()
    doc.add_heading('Marketing', level=2)
    doc.add_paragraph(marketing_text)

    # Add IT section
    doc.add_page_break()
    doc.add_heading('IT', level=2)
    doc.add_paragraph(it_content)

    # Save the document
    doc.save(output_path)

    print(f"CEO Report generated successfully at: {output_path}")

if __name__ == "__main__":
    zip_file_path = 'C:\\Users\\Belay\\IT414_Jschalte\\IT414_NEW\\Week_6\\raw_report_data.zip'
    extract_folder = 'C:\\Users\\Belay\\IT414_Jschalte\\IT414_NEW\\Week_6\\text_files'
    output_path = os.path.join(extract_folder, 'ceo_report.docx')

    extract_raw_data(zip_file_path, extract_folder)
    sales_data = read_sales_data(os.path.join(extract_folder, 'sales.xlsx'))
    operations_content = read_operations_data(os.path.join(extract_folder, 'operations.docx'))
    marketing_text = read_marketing_data(os.path.join(extract_folder, 'marketing.pdf'))
    it_content = read_it_data(os.path.join(extract_folder, 'IT.html'))

    generate_ceo_report(sales_data, operations_content, marketing_text, it_content, output_path)
